#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
სკრიპტი .txt ფაილის .docx ფორმატში გადასაყვანად
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def txt_to_docx(txt_file, docx_file):
    """
    გადააქვს .txt ფაილი .docx ფორმატში ქართული ენის მხარდაჭერით
    """
    # შევქმნათ ახალი Word დოკუმენტი
    doc = Document()
    
    # ფონტის დაყენება
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial Unicode MS'
    font.size = Pt(11)
    
    # წავიკითხოთ .txt ფაილი
    with open(txt_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # თუ ხაზი ცარიელია
        if not line:
            doc.add_paragraph()
            i += 1
            continue
        
        # თუ ხაზი არის გამყოფი (შეიცავს მხოლოდ "=" სიმბოლოებს)
        if line.strip().startswith('=') and all(c == '=' for c in line.strip()):
            i += 1
            continue
        
        # თუ ხაზი იწყება ნომრით და წერტილით (1., 2., 3. და ა.შ.)
        if re.match(r'^\d+\.\s', line):
            # მთავარი სათაური
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(13)
            i += 1
            continue
        
        # თუ ხაზი იწყება ნომრით, წერტილით და ქვენომრით (4.1., 6.2. და ა.შ.)
        if re.match(r'^\d+\.\d+', line):
            # ქვესათაური
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
            i += 1
            continue
        
        # თუ ხაზი იწყება "•" სიმბოლოთი
        if line.strip().startswith('•'):
            # სიის ელემენტი
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            i += 1
            continue
        
        # თუ ხაზი შეიცავს "FUNCTION" ან "END FUNCTION" ან პროგრამირების კონსტრუქციებს
        code_keywords = ['FUNCTION', 'END FUNCTION', 'IF', 'ELSE', 'FOR', 'WHILE', 
                        'RETURN', 'END IF', 'END FOR', 'END WHILE', '//']
        is_code = any(keyword in line for keyword in code_keywords) or line.strip().startswith('//')
        
        if is_code:
            # კოდის სტილი
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            i += 1
            continue
        
        # თუ ხაზი იწყება "-" სიმბოლოთი (მაგრამ არა "--")
        if line.strip().startswith('-') and not line.strip().startswith('--'):
            # სიის ელემენტი
            p = doc.add_paragraph(line.strip(), style='List Bullet')
            i += 1
            continue
        
        # ჩვეულებრივი პარაგრაფი
        # თუ ხაზი მთავრდება არასრულად (არ არის წერტილი, კითხვის ან ძახილის ნიშანი)
        # და შემდეგი ხაზი არ არის ცარიელი, შეერთება
        if i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if next_line and not next_line.startswith(('•', '-', '=', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.')):
                # შევაერთოთ ხაზები
                full_text = line
                i += 1
                while i < len(lines):
                    next_line = lines[i].strip()
                    if not next_line:
                        break
                    if next_line.startswith(('•', '-', '=', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.')):
                        break
                    full_text += ' ' + next_line
                    i += 1
                p = doc.add_paragraph(full_text)
                continue
        
        # ჩვეულებრივი პარაგრაფი
        p = doc.add_paragraph(line)
        i += 1
    
    # დავცალი დოკუმენტი
    doc.save(docx_file)
    print(f"დოკუმენტი შექმნილია: {docx_file}")

if __name__ == "__main__":
    txt_file = "შუალედური-დავალება.txt"
    docx_file = "შუალედური-დავალება.docx"
    txt_to_docx(txt_file, docx_file)
